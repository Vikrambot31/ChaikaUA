from pathlib import Path
import sys

from docx import Document


ROOT = Path(r"C:/ChaikaUA/mobile-app-short")
DEFAULT_SOURCE = ROOT / "ПАНЕЛЬ КОНТРОЛЯ" / "ИНСТРУКЦИЯ_ПО_МОДЕРАЦИИ.md"
DEFAULT_TARGET = ROOT / "ПАНЕЛЬ КОНТРОЛЯ" / "ИНСТРУКЦИЯ_ПО_МОДЕРАЦИИ.docx"


def flush_bullets(document: Document, bullets: list[str]) -> None:
    for bullet in bullets:
        document.add_paragraph(bullet, style="List Bullet")
    bullets.clear()


def main() -> None:
    source = Path(sys.argv[1]) if len(sys.argv) > 1 else DEFAULT_SOURCE
    target = Path(sys.argv[2]) if len(sys.argv) > 2 else DEFAULT_TARGET
    lines = source.read_text(encoding="utf-8").splitlines()
    document = Document()
    bullets: list[str] = []

    for raw_line in lines:
        line = raw_line.rstrip()

        if not line.strip():
            flush_bullets(document, bullets)
            document.add_paragraph("")
            continue

        if line.startswith("# "):
            flush_bullets(document, bullets)
            document.add_heading(line[2:].strip(), level=1)
            continue

        if line.startswith("## "):
            flush_bullets(document, bullets)
            document.add_heading(line[3:].strip(), level=2)
            continue

        if line.startswith("### "):
            flush_bullets(document, bullets)
            document.add_heading(line[4:].strip(), level=3)
            continue

        if line.startswith("- "):
            bullets.append(line[2:].strip())
            continue

        if line[:3].isdigit() and line[1:3] == ". ":
            flush_bullets(document, bullets)
            document.add_paragraph(line)
            continue

        flush_bullets(document, bullets)
        document.add_paragraph(line)

    flush_bullets(document, bullets)
    target.parent.mkdir(parents=True, exist_ok=True)
    document.save(target)


if __name__ == "__main__":
    main()
